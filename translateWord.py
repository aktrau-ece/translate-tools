'''
translateWord.py [2022.11.17]

Function: Batch translates Word files to many languages at once

Requirements:
	Python: https://www.python.org/downloads/
	3 folders (input, intermediary, and output) (in the same folder as this python file) (the intermediary and output folders should be empty)
	Install python-docx library: terminal > "pip install python-docx"
	Install the docx2pdf module: terminal > "pip install docx2pdf"
	Install the easynmt module: terminal > "pip install easynmt" (this might be difficult to get working)
	Pre-trained models for 'opus-mt' (placed in the same folder as this python file), which can be downloaded from https://huggingface.co/models?sort=downloads&search=opus-mt
		The models must also be loaded in the code (see the load_translateModels() function)

Known issues:
	All text within a paragraph is styled the same way (font, colour, bold, italics) (this is because, while differently formatted text is segmented into separate runs, the text will be rejoined prior to translation to increase translation accuracy. After translation, its really hard to know which of the translated text deserves separate formatting.)
	Elements with a static position on page (after conversion to word) will appear in the wrong place (because they dont appear in any 'run.text')
	Translation may not use all CPU cores (I plan to multi-thread this step, so that many translations can be calculated simultaneously)
	Formatting of text within table cells might be fucked up
	Text inside cells inside tables inside other tables are not translated

Limitations:
	Only one *source* language can be set (multiple target languages can be set)
	Rasterized text (text in pictures, videos) isn't translated

Future Steps:
	Multi-thread translation
	Config options verification
	Put the model configs in the cfg namespace
	Translating a paragraph uses the formatting of a selected run based on how much text it has
'''

# MODULES =========================================

try:
	import os # to send operating system commands, and work with os files
	import io
	import shutil # copying files
	import sys
	from types import SimpleNamespace
	import pprint

	from contextlib import contextmanager, redirect_stderr, redirect_stdout
	from os import devnull

	from easynmt import EasyNMT, models # NMT translator
	import docx # python-docx module
	from docx2pdf import convert # to convert word to pdf

	print("Modules imported")
except Exception as e: G.showErr("Error when importing modules", e)

# CONFIG ==========================================

cfg = SimpleNamespace(**{
	# ** Paths can be relative and must end with "\\" **
	"inPath": "InputFolder\\", # folder in which you put your PDF files
	"interPath": "IntermediateFolder\\", # folder for intermediary files created by the script
	"outPath": "OutputFolder\\", # folder in which youll find the translated PDF files once the script is done

	"inLanguage": "en", # input language
	"outLanguage": ["ar"], # list of languages for translated documents (abbreviated form). ex: ["fr", "de", "it", "es", "ar"]

	"convertToPDF": True, # whether or not to: afterwards, convert all translated word documents to PDF
	"verbosity": 5,
	"testingMode": True, # no user input required during runtime, and no error catching

	"translationModels": {
		# translation model locations, and their translation direction in the form: <FOLDER_NAME>: {"inLang": <INPUT_LANGUAGE_ABBREV>, "outLang": <OUTPUT_LANGUAGE_ABBREV>}
		"opus-mt-en-fr\\": {"sourceLang": "en", "targetLang": "fr"},
		"opus-mt-en-de\\": {"sourceLang": "en", "targetLang": "de"},
		"opus-mt-en-it\\": {"sourceLang": "en", "targetLang": "it"},
		"opus-mt-en-es\\": {"sourceLang": "en", "targetLang": "es"},
		"opus-mt-en-ar\\": {"sourceLang": "en", "targetLang": "ar"},
		"opus-mt-en-vi\\": {"sourceLang": "en", "targetLang": "vi"},
		}
	})

# DATA ===========================================

LANGUAGES = { # from: https://py-googletrans.readthedocs.io/en/latest/
	'af': 'afrikaans',
	'sq': 'albanian',
	'am': 'amharic',
	'ar': 'arabic',
	'hy': 'armenian',
	'az': 'azerbaijani',
	'eu': 'basque',
	'be': 'belarusian',
	'bn': 'bengali',
	'bs': 'bosnian',
	'bg': 'bulgarian',
	'ca': 'catalan',
	'ceb': 'cebuano',
	'ny': 'chichewa',
	'zh-cn': 'chinese (simplified)',
	'zh-tw': 'chinese (traditional)',
	'co': 'corsican',
	'hr': 'croatian',
	'cs': 'czech',
	'da': 'danish',
	'nl': 'dutch',
	'en': 'english',
	'eo': 'esperanto',
	'et': 'estonian',
	'tl': 'filipino',
	'fi': 'finnish',
	'fr': 'french',
	'fy': 'frisian',
	'gl': 'galician',
	'ka': 'georgian',
	'de': 'german',
	'el': 'greek',
	'gu': 'gujarati',
	'ht': 'haitian creole',
	'ha': 'hausa',
	'haw': 'hawaiian',
	'iw': 'hebrew',
	'he': 'hebrew',
	'hi': 'hindi',
	'hmn': 'hmong',
	'hu': 'hungarian',
	'is': 'icelandic',
	'ig': 'igbo',
	'id': 'indonesian',
	'ga': 'irish',
	'it': 'italian',
	'ja': 'japanese',
	'jw': 'javanese',
	'kn': 'kannada',
	'kk': 'kazakh',
	'km': 'khmer',
	'ko': 'korean',
	'ku': 'kurdish (kurmanji)',
	'ky': 'kyrgyz',
	'lo': 'lao',
	'la': 'latin',
	'lv': 'latvian',
	'lt': 'lithuanian',
	'lb': 'luxembourgish',
	'mk': 'macedonian',
	'mg': 'malagasy',
	'ms': 'malay',
	'ml': 'malayalam',
	'mt': 'maltese',
	'mi': 'maori',
	'mr': 'marathi',
	'mn': 'mongolian',
	'my': 'myanmar (burmese)',
	'ne': 'nepali',
	'no': 'norwegian',
	'or': 'odia',
	'ps': 'pashto',
	'fa': 'persian',
	'pl': 'polish',
	'pt': 'portuguese',
	'pa': 'punjabi',
	'ro': 'romanian',
	'ru': 'russian',
	'sm': 'samoan',
	'gd': 'scots gaelic',
	'sr': 'serbian',
	'st': 'sesotho',
	'sn': 'shona',
	'sd': 'sindhi',
	'si': 'sinhala',
	'sk': 'slovak',
	'sl': 'slovenian',
	'so': 'somali',
	'es': 'spanish',
	'su': 'sundanese',
	'sw': 'swahili',
	'sv': 'swedish',
	'tg': 'tajik',
	'ta': 'tamil',
	'te': 'telugu',
	'th': 'thai',
	'tr': 'turkish',
	'uk': 'ukrainian',
	'ur': 'urdu',
	'ug': 'uyghur',
	'uz': 'uzbek',
	'vi': 'vietnamese',
	'cy': 'welsh',
	'xh': 'xhosa',
	'yi': 'yiddish',
	'yo': 'yoruba',
	'zu': 'zulu'}

# FUNCTIONS =======================================

class G:
	# Given a directory path (string), returns a list of filenames (strings) (with extensions)
	def listFiles(dir):
		full_list = os.listdir(dir)
		return full_list

	# Returns a files extension, including the dot
	def extension(fileName):
		for ch in range(len(fileName)-1, -1, -1):
			if fileName[ch] == "." :
				return fileName[ch:]
		return ""

	# Returns a file title (filename without the extension)
	def basename(filename):
		for ch in range(len(filename)-1, -1, -1):
			if filename[ch] == "." :
				return filename[:ch]
		return filename

	def wrap(root, affix):
		return affix + root + affix

	# Prints multidimensional arrays nicely by indenting based on layer
	def printArray(array, i = 0): # optional parameter 'i' specifies starting indentation layer. it is incremented as the function recurses
		for s in array:
			if isinstance(s, list): # https://stackoverflow.com/questions/26544091/checking-if-type-list-in-python
				printArray(s, i+1)
			else: print("\t"*i + s) # print as many tabs as is specified by 'i', then the object

	# Similar to printArray, just for dictionaries
	def printDict(dictionary, stripped=False, indentLevel=0):
		if stripped:
			# indentLevel refers to the offset for the display of the entire dictionary (use 0 to align the dictionary with column 0)
			return True
		else:
			# indentLevel refers to the width of each tab (2 is recommended)
			return pprint.pformat(dictionary, sort_dicts=False, indent=indentLevel)

	def showErr(userMsg = "Error", reason = ""):
		if reason == "": print(userMsg)
		else: print(Style.apply(userMsg, "RED") + ":", reason)
		input("Press <ENTER> to exit")
		raise SystemExit

class Translate:
	def translateText(text, sourceLang, targetLang):
		return translateModels.get(sourceLang, targetLang).translate(text, source_lang=sourceLang, target_lang=targetLang, max_new_tokens=512)

	def translateText_robust(text, sourceLang, targetLang):
		if cfg.verbosity >= 4: print("Translating text: " + G.wrap(text, "\""))

		#if "\n" in text: return Translate.translateText_robust(text.replace("\n", " "), sourceLang, targetLang)

		# some characters / combinations of characters will make the translater return some wacky stuff.
		# known cases:
			# strings containing only numbers (and perhaps spaces and periods)
			# the '≈' symbol on its own
			# "NUMBER. (SOME_TEXT)" this will cause the translator to go on about the european council and some other bullshit
			# greek letters, when translated alone, make the translator output nonsense
			# the letter 'r' (by itself) . It makes the translator spit out some garbage. Im going to pre-emptively avoid translating any single letters.
			# "TEXT!)." the translator returns the text, exclamation mark, and closed parenthesis followed by a bunch of periods
		# {
		if (text is None) or (all([(v == " " or v == "\t") for v in text])): res = ""
		elif all([(v.isnumeric() | (v==" ") | (v=="\t") | (v==".")) for v in text]): res = text # if every char in the string is either a number, a tab, a space, or a period
		elif text == "≈": res = text # if the string contains only '≈'

		elif (len(text) > 2) and (text[0].isnumeric()) & (text[1] == "."):
			if cfg.verbosity >= 5: print("Recursing")
			res = text[0:2] + Translate.translateText_robust(text[2:], sourceLang, targetLang)
				
		elif (len(text) == 1): res = text
		
		elif (text.find("!).") != -1) and (all([(v == ".") for v in text[text.find("!).")+2:]])):
			res = (Translate.translateText_robust(text[:text.find("!).")+2], sourceLang, targetLang) + str(text[text.find("!).")+2:]))
		# }

		else:
			res = Translate.translateText(text, sourceLang, targetLang)
			if cfg.verbosity >= 5: print("Clear")

		if (
			len(res) > 3*len(text)
			or (".........." in res)
			or ("----------" in res)
		):
			if cfg.verbosity >= 4: print("Translation looks like some garbage. Using original un-translated text..")
			return text # if the translator bugs and returns a bunch of garbage, return the original untranslated text
		else: return res

	def translateDoc(filename, sourceLang, targetLang):
		doc = docx.Document(cfg.inPath + filename) # Load the word document

		amtParagraph = len(doc.paragraphs)

		''' if cfg.verbosity >= 4:
			print("\t# inline pictures:", len(doc.inline_shapes))
			runList = ""
			for p in doc.paragraphs:
				runList += (str(len(p.runs)) + ", ")
			runList = runList[:len(runList)-2]
			print("\t# paragraphs:", str(len(doc.paragraphs)))
			print("\t# runs in each paragraph:", runList) '''
		
		# translate all paragraphs
		for p in doc.paragraphs:
			if (len([r.text for r in p.runs]) > 0) and any([r.text != "" for r in p.runs]): # if there are any runs in the paragraph, and if any of those run.texts contain characters

				#if cfg.verbosity >= 5: print("Paragraph text: " + G.wrap(p.text, "\""))

				# the issue was that sentences were being split into multiple runs, which means the translator would lack context when translating pieces of sentences (if it were to translate one run at a time)
				# so, the runs were combined so that the translator translates *paragraphs* at a time, providing more accurate translations
				# the problem with this approach is that all the text within a paragraph will be formatted the same, since it will all be put back into one single run
				# we will try to find the most appropriate run (within said paragraph) from which its format (font, italics, etc) is to be used, by: going through the runs until an alphabetical character is found
					# this particularly addresses the case where a paragraph containing bullet points is to be translated; we want to use the format of the first word, and not the format of the bullet point (which is usually some default font like calibri)
				# {
				tempRunText = "".join([r.text for r in p.runs]) # join all of the texts into a single string
				
				for i in range(len(p.runs)): # iterate through runs
					if any([v.isalpha() for v in p.runs[i].text]): # if run.text contains any alphabetical character
						for j in range(len(p.runs)): p.runs[j].text = "" # erase text from all runs within the paragraph

						if not(p.runs[i].font.superscript or p.runs[i].font.subscript): # dont translate superscript or subscript text
							p.runs[i].text = Translate.translateText_robust(tempRunText, sourceLang, targetLang)
						else:
							p.runs[i].text = tempRunText

						continue # exit loop
				# }

		if cfg.verbosity >= 5:
			print("--TABLES--")
			for i in range(len(doc.tables)):
				for j in range(len(doc.tables[i]._cells)):
					for k in range(len(doc.tables[i]._cells[j].paragraphs)):
						for l in range(len(doc.tables[i]._cells[j].paragraphs[k].runs)):
							print(" Table-" + str(i) + " Cell-" + str(j) + " Para-" + str(k) + " Run-" + str(l) + " " + G.wrap(doc.tables[i]._cells[j].paragraphs[k].runs[l].text, "\"") + " Superscript-" + str(doc.tables[i]._cells[j].paragraphs[k].runs[l].font.superscript))


		# translate all tables
		# to preserve text formatting, we navigate down to the 'runs' layer to edit the text
		for i in range(len(doc.tables)):
			for j in range(len(doc.tables[i]._cells)):
				for k in range(len(doc.tables[i]._cells[j].paragraphs)):
					for l in range(len(doc.tables[i]._cells[j].paragraphs[k].runs)):
						tmp_run = doc.tables[i]._cells[j].paragraphs[k].runs[l]

						if not(tmp_run.font.superscript or tmp_run.font.subscript): # dont translate superscript or subscripted text (its probably part of a math equation)
							doc.tables[i]._cells[j].paragraphs[k].runs[l].text = Translate.translateText_robust(doc.tables[i]._cells[j].paragraphs[k].runs[l].text, sourceLang, targetLang) # replace text with translated text

		# Save the document
		filename_out = G.basename(filename) + " -" + targetLang + G.extension(filename)
		doc.save(cfg.interPath + filename_out)

def update_fileList():
	files.inPath = G.listFiles(cfg.inPath)
	files.inPath_docx = [v for v in files.inPath if ((G.extension(v) == ".docx") and (v[0] != "~"))] # temporary files start with '~'
	files.interPath = G.listFiles(cfg.interPath)

@contextmanager
def suppress_stdout_stderr(): # https://gist.github.com/vikjam/755930297430091d8d8df70ac89ea9e2
	with open(devnull, 'w') as fnull:
		with redirect_stderr(fnull) as err, redirect_stdout(fnull) as out:
			yield (err, out)

# DATA TYPES =======================================

class Style():
	BLACK = "\033[0;30m"
	RED = "\033[0;31m"
	GREEN = "\033[0;32m"
	BROWN = "\033[0;33m"
	BLUE = "\033[0;34m"
	PURPLE = "\033[0;35m"
	CYAN = "\033[0;36m"
	LIGHT_GRAY = "\033[0;37m"
	DARK_GRAY = "\033[1;30m"
	LIGHT_RED = "\033[1;31m"
	LIGHT_GREEN = "\033[1;32m"
	YELLOW = "\033[1;33m"
	LIGHT_BLUE = "\033[1;34m"
	LIGHT_PURPLE = "\033[1;35m"
	LIGHT_CYAN = "\033[1;36m"
	LIGHT_WHITE = "\033[1;37m"
	BOLD = "\033[1m"
	FAINT = "\033[2m"
	ITALIC = "\033[3m"
	UNDERLINE = "\033[4m"
	BLINK = "\033[5m"
	NEGATIVE = "\033[7m"
	CROSSED = "\033[9m"
	RESET = "\033[0m"

	def apply(text, colour): return (getattr(Style(), colour.upper()) + text + Style.RESET)

class ModelOrganiser:
	def __init__(self):
		self.repo = {}

	def add(self, model, sourceLang, targetLang):
		self.repo[sourceLang + "-" + targetLang] = model

	def get(self, sourceLang, targetLang):
		if (sourceLang + "-" + targetLang) not in self.repo: # if model is has not been loaded yet (its the first time translating to this language)
			# search for a pre-trained model in 'cfg.translationModels' with the desired translation direction
			for key in cfg.translationModels:
				if cfg.translationModels[key]["sourceLang"] == sourceLang and cfg.translationModels[key]["targetLang"] == targetLang: # if its translation direction is a match
					self.add(EasyNMT(translator=models.AutoModel(key)), sourceLang, targetLang)
					print("Loaded model with translation direction: " + LANGUAGES[sourceLang] + "->" + LANGUAGES[targetLang] + " from: " + key)

			if (sourceLang + "-" + targetLang) not in self.repo: # if desired model was not found in cfg.transationModels..
				G.showErr(reason="No model with translation direction: " + sourceLang + "->" + targetLang + " configured.")
		
		return self.repo[sourceLang + "-" + targetLang]


# MAIN =============================================

# Display config options
def dispConfig():
	print("Config options:")
	print("\n".join([("- " + Style.apply(v, "DARK_GRAY")) for v in G.printDict(vars(cfg), stripped=False, indentLevel=2).split("\n")]))
dispConfig()

# Initialize ModelOrganiser
translateModels = ModelOrganiser()

# generate file list
files = SimpleNamespace(**{})
update_fileList()
print(str(len([v for v in files.inPath if G.extension(v) == ".docx"])) + " Word (.docx) files found in the input folder")
if len(files.interPath) > 0: G.showErr(reason="intermediate folder is not empty.")

# translate all word docs in the input folder
def translateAll():
	for i in range(len(files.inPath_docx)):
		for j in range(len(cfg.outLanguage)):
			print("Translating " + G.wrap(files.inPath_docx[i], "'") + " to " + G.wrap(LANGUAGES[cfg.outLanguage[j]], "'") + "..")
			fileBasename_out = G.basename(files.inPath_docx[i]) + " -" + cfg.outLanguage[j]

			if cfg.testingMode:
				Translate.translateDoc(files.inPath_docx[i], cfg.inLanguage, cfg.outLanguage[j])
			else:
				try: Translate.translateDoc(files.inPath_docx[i], cfg.inLanguage, cfg.outLanguage[j])
				except Exception as e:
					print("Failed to translate " + G.wrap(files.inPath_docx[i], "'") + " to " + G.wrap(LANGUAGES[cfg.outLanguage[j]], "'") + ": " + str(e))
					continue

			if cfg.convertToPDF:
				try:
					with suppress_stdout_stderr(): convert(cfg.interPath + fileBasename_out + ".docx", cfg.outPath)
					os.remove(cfg.interPath + fileBasename_out + ".docx")
				except Exception as e:
						print("Failed to convert " + fileBasename_out + ".docx to PDF: " + str(e))
						continue
			else:
				shutil.move(cfg.interPath + fileBasename_out + ".docx", cfg.outPath + fileBasename_out + ".docx")

			print("Finished translating " + G.wrap(files.inPath_docx[i], "'") + " to " + G.wrap(LANGUAGES[cfg.outLanguage[j]], "'"))

translateAll()

print("End of script:", G.wrap(os.path.basename(__file__), "'"))
input("Press <ENTER> to exit")

# SOURCES ==========================================

'''
Translation and docx
	https://gist.github.com/phillipkent/fcd8276d3984089cddd2f72a52fd00eb
	https://www.geeksforgeeks.org/working-with-paragraphs-in-python-docx-module/
	https://stackabuse.com/text-translation-with-google-translate-api-in-python/
	https://stackoverflow.com/questions/68703369/saving-azure-translated-pdf-files-as-ms-word-docx-files
	https://stackoverflow.com/questions/27691678/finding-image-present-docx-file-using-python/27705408#27705408
	https://www.rikvoorhaar.com/python-docx/
	https://buildmedia.readthedocs.org/media/pdf/python-docx/latest/python-docx.pdf

	https://python-docx.readthedocs.io/en/latest/api/text.html#docx.text.parfmt.ParagraphFormat.keep_together
	https://python-docx.readthedocs.io/en/latest/api/document.html#id1
	https://python-docx.readthedocs.io/en/latest/index.html
	https://python-docx.readthedocs.io/en/latest/api/table.html

Word-to-PDF conversion
	https://fedingo.com/how-to-convert-docx-to-pdf-in-python-linux/

https://stackoverflow.com/questions/4152963/get-name-of-current-script-in-python
'''